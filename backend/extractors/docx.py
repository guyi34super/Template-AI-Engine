# ai-engine/extractors/docx.py
from typing import Dict, Any
from pathlib import Path

def extract_docx(path: str) -> Dict[str, Any]:
    """Extract text from DOCX files using python-docx."""
    try:
        from docx import Document
        doc = Document(path)
        
        # Extract all paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())
        
        # Extract tables if any
        table_blocks = []
        for table in doc.tables:
            headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
            rows = []
            for row in table.rows[1:] if len(table.rows) > 1 else []:
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(row_data)
            
            if headers:
                table_blocks.append({
                    "headers": headers,
                    "rows": rows
                })
        
        # Create text blocks from paragraphs
        text_blocks = [{"page": 1, "text": "\n".join(paragraphs)}] if paragraphs else []
        
        # Extract potential header candidates
        header_candidates = []
        for para in paragraphs[:20]:  # Check first 20 paragraphs
            if len(para) < 100 and any(keyword in para.lower() for keyword in 
                ["title", "name", "date", "total", "amount", "number", "id"]):
                header_candidates.append(para)
        
        # Add table headers as candidates
        for block in table_blocks:
            header_candidates.extend(block["headers"])
        
        # Remove duplicates and limit
        header_candidates = list(dict.fromkeys(header_candidates))[:64]
        
        return {
            "text_blocks": text_blocks,
            "table_blocks": table_blocks,
            "header_candidates": header_candidates
        }
        
    except ImportError:
        # Fallback: return empty structure if python-docx not available
        print("[warn] python-docx not installed. Please install: pip install python-docx")
        return {
            "text_blocks": [{"page": 1, "text": f"Error: Cannot process DOCX file {Path(path).name} - python-docx not installed"}],
            "table_blocks": [],
            "header_candidates": []
        }
    except Exception as e:
        print(f"[warn] DOCX extraction failed: {e}")
        return {
            "text_blocks": [{"page": 1, "text": f"Error processing DOCX file: {str(e)}"}],
            "table_blocks": [],
            "header_candidates": []
        }